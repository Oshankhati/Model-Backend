# rag_engine.py

import os
from pathlib import Path
import fitz  # PyMuPDF
import docx
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings

# ✅ Correct LangChain-compatible wrapper

# embedding_model = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

embedding_model = HuggingFaceEmbeddings(model_name="sentence-transformers/paraphrase-MiniLM-L3-v2")


# embedding_model = HuggingFaceEmbeddings(model_name="all-MiniLM-L12-v2")  # Slightly smaller



# Global vector store
db = None

def extract_text_from_pdf(path):
    doc = fitz.open(path)
    return "\n".join([page.get_text() for page in doc])

def extract_text_from_docx(path):
    doc = docx.Document(path)
    return "\n".join([para.text for para in doc.paragraphs])

def load_documents(folder_path):
    docs = []
    for file in Path(folder_path).glob("*"):
        if file.suffix.lower() == ".pdf":
            text = extract_text_from_pdf(file)
        elif file.suffix.lower() == ".docx":
            text = extract_text_from_docx(file)
        else:
            continue
        docs.append(Document(page_content=text, metadata={"source": file.name}))
    return docs

def chunk_documents(documents):
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    return splitter.split_documents(documents)

def load_and_index_documents(folder_path):
    global db
    print("🔄 Indexing documents...")
    docs = load_documents(folder_path)
    chunks = chunk_documents(docs)
    texts = [chunk.page_content for chunk in chunks]
    metadatas = [chunk.metadata for chunk in chunks]

    db = FAISS.from_texts(texts, embedding_model, metadatas=metadatas)
    print("✅ Indexing complete.")

def search_documents(query, k=4):
    if db is None:
        raise ValueError("Vector DB not initialized.")
    return db.similarity_search(query, k=k)


# ✅ Load only when needed
# def get_embedding_model():
#     return HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
